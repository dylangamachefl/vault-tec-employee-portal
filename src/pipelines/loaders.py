"""
Stage 1: Document loaders.

Each loader returns (raw_text, filename). Loading logic is strictly separated
from cleaning logic — these functions do nothing but extract text.
"""

from pathlib import Path


def load_document(filepath: str) -> tuple[str, str]:
    """
    Dispatch to the appropriate loader based on file extension.

    Returns: (raw_text, filename)
    """
    path = Path(filepath)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _load_pdf(path)
    elif ext == ".docx":
        return _load_docx(path)
    elif ext in (".md", ".txt"):
        return _load_text(path)
    else:
        raise ValueError(f"Unsupported file format: {ext!r} for file {path.name!r}")


# ---------------------------------------------------------------------------
# Private loaders
# ---------------------------------------------------------------------------


def _load_pdf(path: Path) -> tuple[str, str]:
    """Extract text page-by-page from a PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError("pypdf is required for PDF loading. Install it with: uv add pypdf") from e

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    raw_text = "\n\n".join(pages)
    return raw_text, path.name


def _load_docx(path: Path) -> tuple[str, str]:
    """
    Extract text from a DOCX file using python-docx.

    Paragraphs are extracted in document order, joined with double newlines.
    Word tables are rendered as markdown pipe-delimited tables so the
    structure-aware chunker can detect them as atomic protected regions.

    Format for tables:
        | cell1 | cell2 | cell3 |
        |-------|-------|-------|
        | val1  | val2  | val3  |
    """
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX loading. Install it with: uv add python-docx"
        ) from e

    doc = Document(str(path))

    # We need to preserve the document order of paragraphs AND tables.
    # Iterate over the body XML children directly to preserve ordering.
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    blocks: list[str] = []

    for child in doc.element.body:
        if isinstance(child, CT_P):
            # Regular paragraph
            from docx.text.paragraph import Paragraph

            para = Paragraph(child, doc)
            text = para.text.strip()
            if text:
                blocks.append(text)

        elif isinstance(child, CT_Tbl):
            # Word table → render as markdown table
            from docx.table import Table

            table = Table(child, doc)
            rows = table.rows
            if not rows:
                continue

            md_rows: list[str] = []
            for i, row in enumerate(rows):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                # Skip entirely empty rows
                if not any(cells):
                    continue
                md_rows.append("| " + " | ".join(cells) + " |")
                # Insert separator after the first (header) row
                if i == 0:
                    md_rows.append("|" + "|".join(["---"] * len(cells)) + "|")

            if md_rows:
                blocks.append("\n".join(md_rows))

    raw_text = "\n\n".join(blocks)
    return raw_text, path.name


def _load_text(path: Path) -> tuple[str, str]:
    """Read plain text / Markdown files directly."""
    raw_text = path.read_text(encoding="utf-8", errors="replace")
    return raw_text, path.name


# ---------------------------------------------------------------------------
# DOCX heading extraction helper (used by chunker for section title detection)
# ---------------------------------------------------------------------------


def extract_docx_headings(filepath: str) -> list[tuple[int, str]]:
    """
    Return a list of (paragraph_index, heading_text) for every paragraph in the
    DOCX that uses a Heading style.  Used by the chunker for section_title extraction.
    """
    try:
        from docx import Document
    except ImportError:
        return []

    doc = Document(filepath)
    headings: list[tuple[int, str]] = []
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name.startswith("Heading"):
            text = para.text.strip()
            if text:
                headings.append((i, text))
    return headings


def extract_docx_paragraphs_with_style(filepath: str) -> list[tuple[str, str]]:
    """
    Return [(paragraph_text, style_name), ...] for all non-empty paragraphs.
    Used by the chunker to detect headings inline while building text.
    """
    try:
        from docx import Document
    except ImportError:
        return []

    doc = Document(filepath)
    result: list[tuple[str, str]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style_name = para.style.name if para.style else "Normal"
            result.append((text, style_name))
    return result
